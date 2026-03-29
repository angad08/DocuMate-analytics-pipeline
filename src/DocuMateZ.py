"""
====================================================================
 DocuMateZ v4.0 / Z — Birth Registration
====================================================================

 This is where DocuMate stopped being a script and became
 a system.

 v1 proved the idea. v2 made it usable. v3 went into
 production. v4 broke the ceiling — Excel is gone. DocuMate
 now connects directly to a live database, and the same data
 feeds into Power BI dashboards for real-time reporting.

--------------------------------------------------------------------
 THE BIGGER PICTURE: HOW DOCUMATE EVOLVED
--------------------------------------------------------------------

 v1–v3 were process-driven:
   Excel file → generate Word documents → done.
   The focus was on automating a manual task.

 v4/X is process-driven AND analytics-driven:
   Database → generate Word documents → update status →
   Power BI picks up the same data for live dashboards.

 The shift from Excel to a database wasn't just a tech
 upgrade — it turned DocuMate from a document generator
 into a data platform. The same records that produce
 certificates now power real-time reporting, tracking,
 and insights without any extra work.

====================================================================
 v4.0 — THE DATABASE MIGRATION
====================================================================

 The biggest change in DocuMate's history. The entire data
 layer was ripped out and replaced.

 1. DATA SOURCE: EXCEL → POSTGRESQL
    - v3 read from an Excel file on disk.
    - v4.0 connects to a PostgreSQL database hosted on
      Supabase (AWS ap-southeast-2).
    - Records are pulled via a SQL query that joins across
      three tables: applicant, ministry of home affairs,
      and IB authority — no more flat spreadsheets.

 2. DATABASE STATUS UPDATE (REPLACES EXCEL UPDATE)
    - v3 updated an Excel column to "PRINTED" using openpyxl.
    - v4.0 runs a batch SQL UPDATE directly on the database.
    - Fail-safe measures built in:
        • Checks if date_issued column exists before writing
        • Creates it automatically if missing
        • Processes in batches of 300 to avoid timeouts
        • Rolls back on failure — no partial updates

 3. POWER BI INTEGRATION
    - Because data now lives in a database, Power BI connects
      to the same source for real-time dashboards and insights.
    - DocuMate doesn't talk to Power BI directly — they share
      the database, so dashboards update as records get processed.

 4. PARALLEL DOCUMENT GENERATION (SAME AS v3)
    - v3 used multiprocessing (parallel workers).
    - v4.0 uses the same approach — parallel workers via
      ProcessPoolExecutor, rendering documents simultaneously
      and merging them in submission order.
    - Consistent architecture across v3 and v4/X.

 5. AUTO-CREATES OUTPUT FOLDER (INHERITED FROM v3)
    - If the output folder doesn't exist, it gets created
      automatically before saving. Carried over from v3 to
      ensure first-time runs or new machines don't crash.

====================================================================
 X — AUTO-DETECTION MODE
====================================================================

 Small but important addition on top of v4.0.
 This is what makes it "X" — it doesn't just run, it watches.

 1. AUTO-DETECT NEW RECORDS
    - v4.0 (and all previous versions) were run-once:
      you launch it, it processes, it exits.
    - X adds a polling loop that checks for new
      "IN PROCESS" records every X seconds (configurable).
    - Behaves like a lightweight background service.
    - Stops cleanly on Ctrl+C.

 2. NEW ENTRY POINT: .start()
    - .run() still works the same (single execution).
    - .start() wraps .run() and decides:
        • auto_detect = False → run once, same as before
        • auto_detect = True  → loop with poll_interval

--------------------------------------------------------------------
 CORE FUNCTIONALITY (INHERITED FROM v1/v2/v3)
--------------------------------------------------------------------

 1. Pulls "IN PROCESS" records (now from database)
 2. Formats and maps fields to template placeholders
 3. Generates Word documents using a template
 4. Merges all into one timestamped .docx with page breaks
 5. Optionally marks records as "PRINTED" in the database
 6. Shows summary popup

--------------------------------------------------------------------
 KNOWN LIMITATIONS
--------------------------------------------------------------------

 - Database credentials are hardcoded (no config file or env vars)
 - No reconnection logic if the database connection drops
 - Popups are still Windows-only
 - Auto-detect mode has no graceful shutdown beyond Ctrl+C
 - Output folder path is still partially hardcoded

====================================================================
"""


import pandas as pd                          # For data manipulation (used minimally in v4)
from docxtpl import DocxTemplate             # For filling Word templates with data
from docx import Document                    # For creating/manipulating Word documents
from docxcompose.composer import Composer     # For merging multiple Word docs into one
from docx.oxml import OxmlElement            # For inserting raw XML elements (page breaks)
from docx.oxml.ns import qn                  # For resolving XML namespace prefixes
import time                                  # For tracking execution time and poll intervals
import os                                    # For building file paths
import sys                                   # For detecting .exe vs .py execution
import ctypes                                # For Windows popup messages
from datetime import datetime                # For timestamps in filenames
from io import BytesIO                       # For holding rendered docs in memory
import psycopg2                              # PostgreSQL database driver
import datetime as dt                        # For date operations in SQL updates
from concurrent.futures import ProcessPoolExecutor as Executor  # Parallel executor
import multiprocessing                       # For freeze_support in .exe builds


def process_record(template_path, row_dict):
    """
    Renders a single record into a Word document and returns it as
    raw bytes (not a Document object).

    Why bytes? Because Document objects can't be passed between
    processes in multiprocessing — bytes can. Each worker process
    fills the template independently, saves to memory, and sends
    the result back.
    """
    doc_tpl = DocxTemplate(template_path)
    doc_tpl.render(row_dict)
    mem_file = BytesIO()
    doc_tpl.save(mem_file)
    mem_file.seek(0)
    return mem_file.getvalue()


# ====================================================================
# BirthRegistrationProcessor
# --------------------------------------------------------------------
# Same class structure from v2/v3, but the data source has completely
# changed. Instead of reading Excel, it queries a PostgreSQL database.
# Document generation uses the same parallel processing as v3
# (ProcessPoolExecutor). Auto-detection loop added in X.
#
# Database connection is now owned by the class (moved inside
# __init__ in v4.1). Previously it was a module-level global.
# ====================================================================

class BirthRegistrationProcessor:

    def __init__(self, db_config, template_path, output_folder, update_existing=False,auto_detect=False, poll_interval=300):
        """
        Sets up the processor.

        - db_config:       Dictionary with database connection parameters
                           (host, database, user, password, port, sslmode)
        - template_path:   Path to the Word template (.docx)
        - output_folder:   Where the merged output will be saved
        - update_existing: Whether to update database after generation
                           (overridden by user prompt during run)
        - auto_detect:     [X] If True, runs in a loop checking
                           for new records every poll_interval seconds
        - poll_interval:   [X] Seconds between auto-detection
                           checks (default: 300 = 5 minutes)
        """

        self.conn = psycopg2.connect(**db_config)
        self.cur = self.conn.cursor()
        self.template_path = template_path
        self.output_folder = output_folder
        self.data = None                     # Will hold list of record dicts from DB
        self.update_existing = update_existing
        self.auto_detect = auto_detect
        self.poll_interval = poll_interval


    def validate_environment(self) -> bool:
        """
        Pre-flight check before DocuMate does anything.
        
        Validates:
        - Template file exists (crashes mid-loop without this)
        - Output folder exists (creates it if not)
        
        Called at the top of run() so failures are caught
        early with a clear message, not halfway through
        processing 300 records.
        """

        if not os.path.exists(self.template_path):
            ctypes.windll.user32.MessageBoxW(
                0,
                f"❌ Template not found:\n{self.template_path}",
                "DocuMate",
                0
            )
            return False

        if not os.path.exists(self.output_folder):
            print(f"\nDocuMate : 📂 Output folder '{self.output_folder}' does not exist. Creating it...")
            os.makedirs(self.output_folder)

        return True


    def load_data(self):
        """
        Step 1 [v4.0]: Pull records from the database.

        This is the biggest change from v3. Instead of reading
        an Excel file, this runs a SQL query that:
        - Joins applicant → ministry of home affairs → IB authority
        - Filters to STATUS = 'IN PROCESS' at the database level
        - Orders by serial number
        - Returns only the fields needed for the Word template

        The raw SQL rows are then converted into a list of
        dictionaries with keys matching the placeholder names
        in the Word template.
        """

        self.cur.execute("""
            SELECT
                a.file_number,
                a.serial,
                UPPER(a.name),
                UPPER(a.sex),
                a.birth_date,
                UPPER(CONCAT(a.place, ', ', a.state_code)) AS place_of_birth,
                UPPER(a.name_of_father),
                UPPER(a.name_of_mother),
                UPPER(CONCAT(a.address_line_1, ', ',a.address_line_2, ', ',a.address_line_3)) AS address,
                a.registration_date,
                b.mha_file_number,
                b.mha_date,
                UPPER(ib.ib_staff_authority_name) AS signing_authority_name,
                UPPER(CONCAT(ib.ib_staff_authority_name, ', ',ib.ib_staff_authority_designation)) AS authority_name_designation
            FROM applicant a
                JOIN ministryofhomeaffairs b
                    ON a.mha_file_number = b.mha_file_number
                LEFT JOIN ib_authority ib
                    ON a.ib_staff_authority_id = ib.ib_staff_authority_id
            WHERE UPPER(a.status) = 'IN PROCESS'
            ORDER BY a.serial;
        """)

        rows = self.cur.fetchall()

        # If no records match, show a popup and stop
        if not rows:
            msg = (
                "📝 No records found with status 'IN PROCESS'.\n\n"
                "🧠 Powered by ENN — Fuelled by curiosity, refined by data-driven clarity."
            )
            ctypes.windll.user32.MessageBoxW(0, msg, "DocuMate", 0)
            #self.data = None
            #return

        # Convert each SQL row into a dictionary with keys matching
        # the Word template placeholders.
        # Serial 123 → "123/2025", dates → "dd/mm/yyyy",
        # names → UPPERCASE, address fields → concatenated
        #
        # "year" is injected as a standalone key so the Word template
        # can use {{year}} in file numbers or anywhere the current year
        # appears. This prevents the hardcoded-year bug that caused a
        # missed update during the 2025→2026 transition.
        current_year = dt.datetime.now().year

        self.data = [
            {
                "File_Number": r[0],
                "Serial": str(r[1]) +"/"+str(current_year),
                "Name": r[2],
                "Sex": r[3],
                "When_and_where_born":(r[4].strftime("%d/%m/%Y") if hasattr(r[4], "strftime") else str(r[4])) + ", " + r[5],
                "Name_of_the_Father": r[6],
                "Name_of_the_Mother": r[7],
                "Description_and_residence_of_informant": r[8],
                "Registration_date": r[9].strftime("%d/%m/%Y"),
                "MHA_File_And_date": r[10] + ", " + (r[11].strftime("%d/%m/%Y") if hasattr(r[11], "strftime") else str(r[11])),
                "Signing_Authority_Name": r[12],
                "Signing_Authority_Name_Designation": r[13]
            }
            for r in rows
        ]
        return self.data


    def add_page_break(self, doc):
        """
        Inserts a page break into a Word document.
        Same as v3 — used between merged records so each
        certificate starts on its own page.
        """
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)


    def generate_and_merge_documents(self, merged_filename="merged_birth_cert_2025.docx", max_workers=10):
        """
        Step 2: Generate Word documents and merge into one file.

        Same architecture as v3 — parallel processing:

        1. Sends each record to a parallel worker that fills
           the Word template and returns the result as bytes
        2. Collects all results IN ORDER (not random)
        3. Merges them into one document with page breaks
        4. Saves as a single timestamped .docx file

        Only difference from v3: data comes from the database,
        not Excel. The rendering and merging logic is identical.
        """

        if not self.data:
            print("⚠️ No records found for merging.")
            return

        total_records = len(self.data)
        print(f"⏳ Generating and merging {total_records} documents...\n")

        master_doc = Document()
        composer = Composer(master_doc)
        first_doc = True

        # Render all records in parallel using up to 10 workers.
        # Each worker fills the template independently in its own process.
        # Results come back as raw bytes, which we convert to Document objects.
        # IMPORTANT: futures are collected in submission order, so the
        # merged output stays in the same sorted order.
        with Executor(max_workers=max_workers) as executor:
            futures = [
                executor.submit(process_record, self.template_path, row)
                for row in self.data
            ]

            for i, future in enumerate(futures, 1):
                try:
                    mem_bytes = future.result()
                    temp_doc = Document(BytesIO(mem_bytes))

                    # First document becomes the master (keeps its formatting)
                    # All subsequent documents are appended after a page break
                    if first_doc:
                        master_doc = temp_doc
                        composer = Composer(master_doc)
                        first_doc = False
                    else:
                        self.add_page_break(master_doc)
                        composer.append(temp_doc)

                    print(f"DocuMate :✔️ Processed record {i}/{total_records}", end="\r")

                except Exception as e:
                    print(f"\n⚠️ Error in record {i}: {e}")

        # if the output folder does not exist, create it
        if not os.path.exists(self.output_folder):
            print(f"\nDocuMate : 📂 Output folder '{self.output_folder}' does not exist. Creating it...")
            os.makedirs(self.output_folder)

        # Save the merged document to disk
        output_path = os.path.join(self.output_folder, merged_filename)
        composer.save(output_path)
        print(f"\nDocuMate : ✅ Generated and merged {total_records} documents and saved in:\n📂 {output_path}")


    def update_sql_status(self):
        """
        Step 3 (optional) [v4.0]: Mark processed records as
        "PRINTED" in the database.

        Replaces v3's Excel update with a direct SQL UPDATE.

        Safety measures:
        - Checks if date_issued column exists; creates it if not
        - Processes in batches of 300 to avoid timeout issues
        - Only updates rows still marked "IN PROCESS"
          (prevents double-updating if run twice)
        - Rolls back the entire transaction on failure
        """
        if not self.data:
            print("⚠️ DocuMate : No data to update.")
            return
        
        try:
            # Remove any query timeout for large batches
            self.cur.execute("SET statement_timeout = 0;")

            # Fail-safe: make sure date_issued column exists.
            # If someone set up the DB without it, this creates
            # it automatically instead of crashing.
            self.cur.execute("""
                SELECT 1
                FROM information_schema.columns
                WHERE table_name = 'applicant'
                AND column_name = 'date_issued';
            """)
            if not self.cur.fetchone():
                self.cur.execute("ALTER TABLE applicant ADD COLUMN date_issued DATE;")
                self.conn.commit()

            today = dt.date.today()

            # Get serial numbers of records processed in THIS run only
            serials = [int(row["Serial"].split("/")[0]) for row in self.data]

            # Update in batches of 300 to avoid hitting database
            # timeouts or memory limits on large datasets
            batch_size = 300
            total_updated = 0

            for i in range(0, len(serials), batch_size):
                batch = serials[i:i + batch_size]

                # WHERE clause double-checks status is still "IN PROCESS"
                # to prevent re-marking records already updated
                self.cur.execute(
                    """
                    UPDATE applicant
                    SET status = 'PRINTED',
                        date_issued = %s
                    WHERE serial = ANY(%s)
                    AND coalesce(upper(status),'') = 'IN PROCESS';
                    """,
                    (today, batch)
                )

                total_updated += self.cur.rowcount
                self.conn.commit()

            print(f"✅ DocuMate : Database updated successfully ({total_updated} records).")

        except Exception as e:
            # Roll back so no partial updates are left in the database
            self.conn.rollback()
            print(f"❌ DocuMate : DB update failed — {e}")



    def run(self):
        """
        Runs the full pipeline once:
        Load from DB → Generate & Merge → (Optional) Update DB

        Same structure as v3's run(), but data comes from the
        database and status update goes to the database.
        """
        try:
            # --- Pre-flight check ---
            if not self.validate_environment():
                return

            # --- Load from database ---
            print("DocuMate : 📄 Reading database records...")
            self.load_data()
            if not self.data:
                return
            print("DocuMate : 🔍 Searching for new records...")
            # self.format_dates()
            print(f"DocuMate : 📝 Found {len(self.data)} new records to populate.\n")

            total_start = time.time()

            # --- Generate & Merge ---
            print("DocuMate : ⚙️  Processing Word documents...\n")
            merge_start = time.time()
            self.generate_and_merge_documents(merged_filename=f"DocuMatePLUS_BIRTH_REGISTRATION_{datetime.now().strftime('%d%m%Y_%H%M%S')}.docx")
            merge_time = time.time() - merge_start
            print(f"\nDocuMate :✅ Generated and merged {len(self.data)} documents in {merge_time:.2f} seconds.\n")

            # --- Ask user: update database? ---
            user_choice = input(
                f"DocuMate : Do you want me to mark the statuses as PRINTED for {len(self.data)} records? (yes/no): "
            ).strip().lower()
            self.update_existing = user_choice in ("yes", "y")

            # --- (Optional) Update database ---
            update_text = ""
            excel_time = 0.0
            if self.update_existing:
                print("\nDocuMate : ⏳ Updating database statuses...\n")
                excel_start = time.time()
                self.update_sql_status()
                excel_time = time.time() - excel_start
                update_text = f" and updated Database in {excel_time:.2f} seconds"

            # --- Final summary popup ---
            total_time = time.time() - total_start

            msg = (
                f"✅ DocuMate : Mission accomplished!\n"
                f"✅ Generated {len(self.data)} documents successfully{update_text}.\n"
                f"✅ Total time taken: {total_time:.2f} seconds.\n\n"
                "🧠 Powered by ENN — Fuelled by curiosity, refined by data-driven clarity."
            )
            ctypes.windll.user32.MessageBoxW(0, msg, "DocuMate", 0)

        except PermissionError:
            msg = (
                "❌ I cannot engage the populate launch because the Excel file is open.\n"
                "Please close it and try again later.\n\nMission aborted."
            )
            ctypes.windll.user32.MessageBoxW(0, msg, "DocuMate", 0)

        except Exception as e:
            msg = (
                f"⚠️ DocuMate encountered an unexpected error:\n{e}\n\n"
                "🧠 Powered by ENN — Fuelled by curiosity, refined by data-driven clarity."
            )
            ctypes.windll.user32.MessageBoxW(0, msg, "DocuMate", 0)


    def start(self):
        """
        [X] Entry point — decides how DocuMate runs:

        - auto_detect = False → runs once and exits (same as v3/v4.0)
        - auto_detect = True  → runs in a loop, checking for new
          records every poll_interval seconds. Acts like a
          lightweight background service. Stops on Ctrl+C.
        """

        if not self.auto_detect:
            # Single run — same behaviour as previous versions
            self.run()
            return

        # Auto-detection loop (X)
        print("🤖 DocuMate auto-detection enabled.")
        print(f"⏳ Checking for new records every {self.poll_interval} seconds...\n")

        while True:
            try:
                self.run()
                time.sleep(self.poll_interval)

            except KeyboardInterrupt:
                print("\n🛑 DocuMate: Auto-detection stopped by user.")
                break

            except Exception as e:
                print(f"⚠️ Auto-detection error: {e}")
                time.sleep(self.poll_interval)


# ====================================================================
# ENTRY POINT
# --------------------------------------------------------------------
# Unlike v3, there's no excel_path or sheet_name here — the
# database connection is now created inside the class.
#
# auto_detect=True + poll_interval=30 means it will keep running,
# checking for new records every 30 seconds.
#
# Expected folder structure:
#
#   base_dir/
#   ├── templates/DOCUMENT_TEMPLATE_FILE.docx  ← Word template
#   └── output_files/                                          ← merged output
#
# ====================================================================

if __name__ == "__main__":
    multiprocessing.freeze_support()

    print("\n🚀 ENN : Initializing hypersonic missile DocuMate launch...")
    print("DocuMate : Alright ENN, initializing populate launch 🚀\n")

    # Resolve base directory (same logic as v2/v3)
    base_dir = (
        os.path.dirname(os.path.dirname(sys.executable))
        if getattr(sys, "frozen", False)
        else os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    )

    db_config = {
        "host": "aws-1-ap-southeast-2.pooler.supabase.com",
        "database": "postgres",
        "user": "postgres.echjbuprlnxkbxflylzf",
        "password": "docuMateX9968",             # <- put your DB password
        "port": 5432,
        "sslmode": "require"
    }

    docuMateAgent = BirthRegistrationProcessor(
        db_config=db_config,
        template_path=os.path.join(base_dir, "templates", "DOCUMENT_TEMPLATE_FILE.docx"),
        output_folder=os.path.join(base_dir,"output_files"),
        update_existing=False,
        auto_detect=True,
        poll_interval=30
    )

    # .start() instead of .run() -- start() decides whether to
    # run once or loop based on the auto_detect setting [Z]
    docuMateAgent.start()
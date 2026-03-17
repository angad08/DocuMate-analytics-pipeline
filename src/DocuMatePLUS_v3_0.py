"""
====================================================================
 DocuMatePLUS v3.0 — Birth Registration
====================================================================

 This is the version that went into production.

 v1 proved the idea. v2 made it usable. v3 made it real —
 deployed as a standalone .exe with an icon, used by the
 Indian Consulate for birth registration certificate
 generation starting 2025.

--------------------------------------------------------------------
 WHAT'S NEW IN v3 (vs v2)
--------------------------------------------------------------------

 1. MERGED OUTPUT (SINGLE FILE)
    - v2 created one .docx per record (100 records = 100 files).
    - v3 merges all records into ONE document with page breaks,
      sorted by Serial number, saved with a timestamped filename.
      Much easier to print and handle.

 2. PARALLEL PROCESSING (MULTIPROCESSING)
    - v2 generated documents one by one, sequentially.
    - v3 uses parallel workers (up to 10) to render documents
      simultaneously, then merges them in order.
      Significantly faster for large batches.

 3. DATA VALIDATION BEFORE PROCESSING
    - v2 only checked if records exist.
    - v3 checks for:
        • Missing 'STATUS' column entirely
        • Blank/missing values in mandatory fields
        • Duplicate File_Number entries
      Stops with a clear message if anything is wrong,
      instead of generating broken documents.

 4. EXCEL STATUS UPDATE AFTER PROCESSING
    - v2 had no way to mark records as done.
    - v3 asks the user after generation:
      "Do you want to mark these as PRINTED?"
      If yes, updates the STATUS column to "PRINTED" and
      stamps the Date_Printed — while preserving existing
      formulas and formatting in the Excel file.

 5. PACKAGED AS STANDALONE .EXE WITH ICON
    - Built with PyInstaller as a single executable.
    - Non-technical users just double-click to run.
    - No Python, no terminal knowledge needed.

--------------------------------------------------------------------
 CORE FUNCTIONALITY (INHERITED FROM v1/v2)
--------------------------------------------------------------------

 1. Reads data from DocuMate_Records/DocuMate_DataFrame.xlsx → DocuMateSRC
 2. Filters rows where STATUS == "IN PROCESS"
 3. Formats date fields to dd/mm/yyyy
 4. Generates Word documents using a template
 5. Displays summary (records processed + time taken)

--------------------------------------------------------------------
 KNOWN LIMITATIONS
--------------------------------------------------------------------

 - Popups are Windows-only (ctypes.windll)
 - Output folder path is still partially hardcoded in __main__
 - No GUI — runs in a console window with popup alerts
 - No rollback if Excel update fails midway

====================================================================

--------------------------------------------------------------------
 IDEA EXPLORED BUT DROPPED: AUTOMATED SIGNATURES
--------------------------------------------------------------------

 During v3 development, automating the signing authority's
 signature and name stamp on each certificate was considered.
 Technically feasible — would have been a game changer for
 end-to-end automation.

 It was dropped for two reasons:

 1. PRINT QUALITY
    The operator running DocuMate had a black-and-white
    printer. A signature reproduced in B/W would not look
    genuine on a dispatched certificate and could be
    questioned in the real world.

 2. ETHICS AND LEGALITY
    The signing authority officer personally signs above
    their stamped name on each certificate. Even automating
    just the stamp and printed name (without the signature)
    was considered, but was stopped on ethical and legal
    grounds — automating an officer's identity mark on
    official government documents crosses a line that
    shouldn't be crossed without explicit authorization.

====================================================================
"""


from __future__ import annotations

import multiprocessing                       # For parallel document generation
import pandas as pd                          # For reading and manipulating Excel data
from docxtpl import DocxTemplate             # For filling Word templates with data
from docx import Document                    # For creating/manipulating Word documents
from docxcompose.composer import Composer     # For merging multiple Word docs into one
from docx.oxml import OxmlElement            # For inserting raw XML elements (page breaks)
from docx.oxml.ns import qn                  # For resolving XML namespace prefixes
import time                                  # For tracking execution time
import os                                    # For building file paths
import sys                                   # For detecting .exe vs .py execution
import ctypes                                # For Windows popup messages
from openpyxl import load_workbook           # For updating Excel without destroying formulas
from datetime import datetime                # For timestamps in filenames and date stamps
from io import BytesIO                       # For holding rendered docs in memory (no temp files)
from concurrent.futures import ProcessPoolExecutor as Executor, Future
from typing import Optional, Dict, List, Any
from pandas import DataFrame


# ====================================================================
# PyInstaller build command (for reference):
# pyinstaller --onefile --icon="path/to/documate.ico" --name DocuMatePLUS_BirthRegistration DocuMateBirthRegistrationPLUS.py
# ====================================================================


def process_record(template_path: str, row_dict: Dict[str, Any]) -> bytes:
    """
    Renders a single record into a Word document and returns it as
    raw bytes (not a Document object).

    Why bytes? Because Document objects can't be passed between
    processes in multiprocessing — bytes can. Each worker process
    fills the template independently, saves to memory, and sends
    the result back.
    """
    doc_tpl: DocxTemplate = DocxTemplate(template_path)
    doc_tpl.render(row_dict)
    mem_file: BytesIO = BytesIO()
    doc_tpl.save(mem_file)
    mem_file.seek(0)
    return mem_file.getvalue()


# ====================================================================
# BirthRegistrationProcessor
# --------------------------------------------------------------------
# Same class from v2, but with major upgrades:
# - Smarter filtering with validation checks
# - Merged document output instead of individual files
# - Parallel processing for speed
# - Optional Excel STATUS update after generation
# ====================================================================

class BirthRegistrationProcessor:

    excel_path: str
    sheet_name: str
    template_path: str
    output_folder: str
    data: Optional[DataFrame]
    update_existing: bool

    def __init__(
        self,
        excel_path: str,
        sheet_name: str,
        template_path: str,
        output_folder: str,
        update_existing: bool = False
    ) -> None:
        """
        Sets up all paths and settings.

        - excel_path:      Path to the Excel file with records
        - sheet_name:      Which sheet to read (e.g. "Sheet2")
        - template_path:   Path to the Word template (.docx)
        - output_folder:   Where the merged output will be saved
        - update_existing: Whether to update Excel after generation
                           (overridden by user prompt during run)
        """
        self.excel_path: str = excel_path
        self.sheet_name: str = sheet_name
        self.template_path: str = template_path
        self.output_folder: str = output_folder
        self.data: Optional[DataFrame] = None
        self.update_existing: bool = update_existing


    def load_data(self) -> None:
        """
        Step 1: Read the Excel file into memory.
        Same as v2 — loads all rows from the specified sheet.
        """
        self.data = pd.read_excel(self.excel_path, sheet_name=self.sheet_name)


    def filter_records(self) -> None:
        """
        Step 2: Filter and validate records.

        This is the big upgrade from v2. Instead of just filtering
        by STATUS, v3 runs a gauntlet of checks before allowing
        processing to continue:

        Check 1 → Is there any data at all?
        Check 2 → Does the 'STATUS' column exist?
        Check 3 → Are there any "IN PROCESS" records?
        Check 4 → Do all mandatory fields have values?
        Check 5 → Are there any duplicate File_Numbers?

        If ANY check fails, processing stops with a clear message.
        self.data is set to None so downstream methods know to skip.
        """

        # --- Check 1: Is there any data at all? ---
        if self.data is None or self.data.empty:
            msg: str = (
                "📝 Found no records to populate.\n\n"
                "🧠 Powered by ENN — Fuelled by curiosity, refined by data-driven clarity."
            )
            ctypes.windll.user32.MessageBoxW(0, msg, "DocuMate", 0)
            self.data = None
            return

        # --- Check 2: Does the 'STATUS' column exist? ---
        if "STATUS" not in self.data.columns:
            msg: str = (
                "⚠️ 'STATUS' column not found. Please check the Excel file.\n\n"
                "🧠 Powered by ENN — Fuelled by curiosity, refined by data-driven clarity."
            )
            ctypes.windll.user32.MessageBoxW(0, msg, "DocuMate", 0)
            self.data = None
            return

        # --- Check 3: Filter to "IN PROCESS" only ---
        self.data = self.data[self.data["STATUS"].str.upper() == "IN PROCESS"]

        if self.data.empty:
            msg: str = (
                "📝 No records found with STATUS 'In Process'.\n\n"
                "🧠 Powered by ENN — Fuelled by curiosity, refined by data-driven clarity."
            )
            ctypes.windll.user32.MessageBoxW(0, msg, "DocuMate", 0)
            self.data = None
            return

        # --- Check 4: Are any mandatory fields blank or missing? ---
        # Skips columns that are allowed to be empty
        # (STATUS, Date_Printed, Date_Issued are not mandatory)
        cols_to_check: List[str] = [col for col in self.data.columns if col not in ["STATUS", "Date_Printed", "Date_Issued"]]
        missing_mask: DataFrame = (
            self.data[cols_to_check].isna()
            | self.data[cols_to_check].apply(lambda col: col.astype(str).str.strip() == "")
        )

        rows_with_missing: DataFrame = self.data[missing_mask.any(axis=1)]

        if not rows_with_missing.empty:
            # Print which rows have problems and which fields are missing
            print("\nDocuMate :⚠️ Records with missing mandatory fields detected:\n")

            for idx, row in rows_with_missing.iterrows():
                missing_cols: List[str] = [
                    col for col in cols_to_check
                    if pd.isna(row[col]) or str(row[col]).strip() == ""
                ]
                # +2 because: row 1 = header, row index starts at 0
                print(f"➡️ Row {idx + 2} | Missing fields: {', '.join(missing_cols)}")

            msg: str = (
                f"⚠️ Warning: Found {len(rows_with_missing)} record(s) with missing values.\n"
                "Please correct the highlighted rows in the Excel file before proceeding.\n\n"
                "🧠 Powered by ENN — Fuelled by curiosity, refined by data-driven clarity."
            )
            ctypes.windll.user32.MessageBoxW(0, msg, "DocuMate", 0)
            self.data = None
            return

        # --- Check 5: Are there duplicate File_Numbers? ---
        # Duplicates could mean the same person was entered twice,
        # which would produce duplicate certificates
        duplicates: DataFrame = self.data[self.data.duplicated(subset=["File_Number"], keep=False)]
        if not duplicates.empty:
            print(duplicates[["File_Number","Serial","Name","Date_Issued"]])
            msg: str = (
                f"⚠️ Warning: Found {len(duplicates)} duplicate File_Number entries in 'In Process' records:\n"
                "Please check the Excel file for duplicates.\n\n"
                "🧠 Powered by ENN — Fuelled by curiosity, refined by data-driven clarity."
            )
            ctypes.windll.user32.MessageBoxW(0, msg, "DocuMate", 0)
            self.data = None
            return


    def format_dates(self, date_columns: List[str]) -> None:
        """
        Step 3: Clean up date columns.
        Same as v2 — converts to dd/mm/yyyy for clean document output.
        """
        for col in date_columns:
            if col in self.data.columns:
                self.data[col] = pd.to_datetime(self.data[col], errors="coerce").dt.strftime("%d/%m/%Y")


    def add_page_break(self, doc: Document) -> None:
        """
        Inserts a page break into a Word document.
        Used between merged records so each certificate starts
        on its own page when printed.
        """
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        br: OxmlElement = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)


    def generate_and_merge_documents(
        self,
        merged_filename: str = f"DocuMatePLUS_BIRTH_REGISTRATION_{datetime.now().strftime('%d%m%Y')}.docx",
        max_workers: int = 10
    ) -> None:
        """
        Step 4: The main event — generate and merge all documents.

        How it works:
        1. Sorts records by Serial number (ascending)
        2. Sends each record to a parallel worker that fills
           the Word template and returns the result as bytes
        3. Collects all results IN ORDER (not random)
        4. Merges them into one document with page breaks
        5. Saves as a single timestamped .docx file

        Output example:
        "DocuMatePLUS_BIRTH_REGISTRATION_10022026.docx"
        """
        if self.data is None or self.data.empty:
            print("DocuMate : ⚠️ No records found for merging.")
            return

        # Sort by the numeric part of Serial (e.g. "123/2025" → sort by 123)
        # so the merged document is in a logical order for printing
        if "Serial" in self.data.columns:
            try:
                self.data["__serial_num__"] = (
                    self.data["Serial"].astype(str).str.extract(r"(\d+)", expand=False).fillna("0").astype(int)
                )
                self.data = self.data.sort_values("__serial_num__")
            except Exception:
                print("⚠️ Some Serial values could not be converted, sorting skipped.")

        total_records: int = len(self.data)
        print(f"DocuMate : ⏳ Generating and merging {total_records} documents...\n")

        # Start with an empty master document that others will be merged into
        master_doc: Document = Document()
        composer: Composer = Composer(master_doc)
        first_doc: bool = True

        # Render all records in parallel using up to 10 workers.
        # Each worker fills the template independently in its own process.
        # Results come back as raw bytes, which we convert to Document objects.
        # IMPORTANT: futures are collected in submission order, so the
        # merged output stays in the same sorted order.
        # Inject current year into each record's context so the Word
        # template can use {{year}} wherever the year appears (e.g. file numbers).
        # This avoids hardcoding the year in the template, which caused a missed
        # update during the 2025→2026 transition. One dict key addition per
        # record costs nanoseconds — no performance impact.
        current_year: int = datetime.now().year

        with Executor(max_workers=max_workers) as executor:
            futures: List[Future[bytes]] = []
            for row in self.data.itertuples(index=False):
                row_dict: Dict[str, Any] = row._asdict()
                row_dict["year"] = current_year
                futures.append(
                    executor.submit(process_record, self.template_path, row_dict)
                )

            for i, future in enumerate(futures, 1):
                try:
                    mem_bytes: bytes = future.result()
                    temp_doc: Document = Document(BytesIO(mem_bytes))

                    # First document becomes the master (keeps its formatting)
                    # All subsequent documents are appended after a page break
                    if first_doc:
                        master_doc = temp_doc
                        composer = Composer(master_doc)
                        first_doc = False
                    else:
                        self.add_page_break(master_doc)
                        composer.append(temp_doc)

                    print(f"DocuMate :✔️ Processed record {i}/{total_records}", end="\r")

                except Exception as e:
                    print(f"\n⚠️ Error in record {i}: {e}")

        # if the output folder does not exist, create it
        if not os.path.exists(self.output_folder):
            print(f"\nDocuMate : 📂 Output folder '{self.output_folder}' does not exist. Creating it...")
            os.makedirs(self.output_folder)
            
        # Save the merged document
        output_path: str = os.path.join(self.output_folder, merged_filename)
        composer.save(output_path)
        print(f"\nDocuMate : ✅ Generated and merged {total_records} documents (in Serial order) and saved in:\n📂 {output_path}")

        # Clean up the temporary sorting column
        if "__serial_num__" in self.data.columns:
            self.data.drop(columns=["__serial_num__"], inplace=True)


    def update_excel_STATUS(self) -> None:
        """
        Step 5 (optional): Mark processed records as "PRINTED" in Excel.

        Uses openpyxl (not pandas) to update the Excel file directly.
        This is important because pandas would destroy any existing
        formulas, formatting, or conditional formatting in the sheet.

        What it does:
        - Opens the workbook with openpyxl (preserves everything)
        - Finds all rows where STATUS == "IN PROCESS"
        - Changes STATUS to "PRINTED"
        - Stamps today's date in the Date_Printed column
          (creates the column if it doesn't exist)
        - Saves the workbook
        """
        try:
            wb = load_workbook(self.excel_path)
            ws = wb[self.sheet_name]

            # Build a map of column names → column numbers from the header row
            headers: Dict[str, int] = {cell.value: idx for idx, cell in enumerate(ws[1], 1)}
            if "STATUS" not in headers:
                print("⚠️ STATUS column not found — skipping Excel update.")
                return

            # If Date_Printed column doesn't exist yet, create it
            if "Date_Printed" not in headers:
                next_col: int = ws.max_column + 1
                ws.cell(row=1, column=next_col, value="Date_Printed")
                headers["Date_Printed"] = next_col

            today: str = datetime.now().strftime("%d/%m/%Y")

            # Find all "IN PROCESS" rows first, then update in bulk
            # (avoids scanning the sheet multiple times)
            STATUS_col: int = headers["STATUS"]
            date_col: int = headers["Date_Printed"]
            in_process_rows: List[int] = [
                r for r in range(2, ws.max_row + 1)
                if str(ws.cell(row=r, column=STATUS_col).value).strip().upper() == "IN PROCESS"
            ]

            for r in in_process_rows:
                ws.cell(row=r, column=STATUS_col, value="PRINTED")
                ws.cell(row=r, column=date_col, value=today)

            wb.save(self.excel_path)
            print(f"🗂️ DocuMate : Sheet '{self.sheet_name}' updated successfully on {today} at {time.strftime('%H:%M:%S')}.")

        except PermissionError:
            print("❌ Cannot update Excel while file is open. Please close it and retry.")
        except Exception as e:
            print(f"⚠️ DocuMate : Error updating Excel — {e}")


    def run(self) -> None:
        """
        Runs the full pipeline in order:
        Load → Validate & Filter → Format → Generate & Merge → (Optional) Update Excel

        Same structure as v2's run(), but with:
        - More validation before processing starts
        - Merged output instead of individual files
        - User prompt for Excel STATUS update at the end
        """
        print("Deploying DocuMate...\n")
        print("DocuMate : Alright ENN, Ready to engage - initializing populate launch 🚀\n")
        try:
            # --- Load ---
            print("DocuMate : 📄 Reading Table...")
            self.load_data()

            # --- Validate & Filter ---
            print("DocuMate : 🔍 Searching for new applicants...")
            self.filter_records()
            if self.data is None:
                return  # Validation failed — message already shown

            # --- Format dates ---
            self.format_dates(["Registration_date"])
            print(f"DocuMate : 📝 Found {len(self.data)} new creations of Lord Bramha 🙏🏻...")

            total_start: float = time.time()

            # --- Generate & Merge ---
            print("DocuMate : ⚙️ Processing Word documents...")
            merge_start: float = time.time()
            self.generate_and_merge_documents()
            merge_time: float = time.time() - merge_start
            print(f"\nDocuMate :✅ Generated and merged {len(self.data)} documents in {merge_time:.2f} seconds.\n")

            # --- Ask user: update Excel? ---
            # This prompt runs in the console. User types yes/no.
            user_choice: str = input(
                f"DocuMate : Do you want me to mark the STATUSes as PRINTED for {len(self.data)} records? (yes/no): "
            ).strip().lower()
            self.update_existing = user_choice in ("yes", "y")

            # --- (Optional) Update Excel ---
            update_text: str = ""
            excel_time: float = 0.0
            if self.update_existing:
                print("\nDocuMate : ⏳ Updating Excel STATUSes...\n")
                excel_start: float = time.time()
                self.update_excel_STATUS()
                excel_time = time.time() - excel_start
                update_text = f" and updated Excel in {excel_time:.2f} seconds"

            # --- Final summary popup ---
            total_time: float = time.time() - total_start

            msg: str = (
                f"✅ DocuMate : Mission accomplished!\n"
                f"✅ Generated {len(self.data)} documents successfully{update_text}.\n"
                f"✅ Total time taken: {total_time:.2f} seconds.\n\n"
                "🧠 Powered by ENN — Fuelled by curiosity, refined by data-driven clarity."
            )
            ctypes.windll.user32.MessageBoxW(0, msg, "DocuMate", 0)

        except PermissionError:
            # Excel file is locked by another program
            msg: str = (
                "❌ I cannot engage the populate launch because the Excel file is open.\n"
                "Please close it and try again later.\n\nMission aborted."
            )
            ctypes.windll.user32.MessageBoxW(0, msg, "DocuMate", 0)

        except Exception as e:
            # Catch-all for anything unexpected
            msg: str = (
                f"⚠️ DocuMate encountered an unexpected error:\n{e}\n\n"
                "🧠 Powered by ENN — Fuelled by curiosity, refined by data-driven clarity."
            )
            ctypes.windll.user32.MessageBoxW(0, msg, "DocuMate", 0)


# ====================================================================
# ENTRY POINT
# --------------------------------------------------------------------
# freeze_support() is required for multiprocessing to work correctly
# when packaged as a PyInstaller .exe. Without it, the .exe would
# spawn infinite child processes on Windows.
#
# Expected folder structure:
#
#   base_dir/
#   ├── DocuMate_Records/DocuMate_DataFrame.xlsx     ← input data
#   ├── DocuMate_Templates/DOCUMENT_TEMPLATE_FILE.docx  ← Word template
#   └── FILES/                                       ← merged output goes here
#
# ====================================================================

if __name__ == "__main__":
    multiprocessing.freeze_support()

    # Resolve base directory (same logic as v2)
    base_dir: str = (
        os.path.dirname(os.path.dirname(sys.executable))
        if getattr(sys, "frozen", False)
        else os.path.dirname(__file__)
    )

    docuMatePLUSAgent: BirthRegistrationProcessor = BirthRegistrationProcessor(
        excel_path=os.path.join(base_dir, "DocuMate_Records", "DocuMate_DataFrame.xlsx"),
        sheet_name="DocuMateSRC",
        template_path=os.path.join(base_dir, "DocuMate_Templates", "DOCUMENT_TEMPLATE_FILE.docx"),
        output_folder=os.path.join(base_dir,"FILES"),
        update_existing=False,
    )

    docuMatePLUSAgent.run()